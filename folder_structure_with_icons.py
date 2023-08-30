
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime

# Define the icons for different file types (here using placeholders, you can replace them with actual icons)
file_icons = {
    '.py': "[Python Icon]",
    '.js': "[JS Icon]",
    '.html': "[HTML Icon]",
    '.jpg': "[JPG Icon]",
    '.png': "[PNG Icon]",
    '.mp4': "[MP4 Icon]",
    '.mp3': "[MP3 Icon]",
    '.yaml': "[YAML Icon]",
    '.json': "[JSON Icon]",
    '.go': "[Go Icon]",
    # Add more file types and their icons here
}

def list_files(startpath, ignore_git, use_icons, doc=None):
    folder_icon = "📁" if use_icons else "Folder:"
    file_icon = "📄" if use_icons else "File:"

    structure = ""
    
    for root, dirs, files in os.walk(startpath):
        if ignore_git and ".git" in root:
            continue
            
        level = root.replace(startpath, '').count(os.sep)
        indent = ' ' * 4 * level

        structure += f"{indent}{folder_icon} {os.path.basename(root)}\n"
        if doc:
            p = doc.add_paragraph()
            run = p.add_run(f"{indent}{folder_icon} {os.path.basename(root)}")
            font = run.font
            font.size = Pt(12)

        sub_indent = ' ' * 4 * (level + 1)

        for f in files:
            file_ext = os.path.splitext(f)[1]
            specific_file_icon = file_icons.get(file_ext, file_icon)  # Fall back to generic file icon
            structure += f"{sub_indent}{specific_file_icon} {f}\n"
            if doc:
                p = doc.add_paragraph()
                run = p.add_run(f"{sub_indent}{specific_file_icon} {f}")
                font = run.font
                font.size = Pt(12)
                
    return structure

def main():
    folder_path = input("Enter the path of the folder (hit Enter for current folder): ")
    if not folder_path:
        folder_path = os.getcwd()

    output_docx = input("Output to .docx file? (Y/N): ").strip().lower()
    include_git = input("Include .git files? (Y/N): ").strip().lower()
    use_icons = input("Use icons for folders and files? (Y/N): ").strip().lower()

    ignore_git = include_git != 'y'
    use_icons = use_icons == 'y'

    folder_structure = list_files(folder_path, ignore_git, use_icons)

    if output_docx == 'y':
        doc = Document()
        list_files(folder_path, ignore_git, use_icons, doc)

        docx_filename = "{}_folder_structure.docx".format(os.path.basename(folder_path))

        if os.path.exists(docx_filename):
            overwrite = input(f"{docx_filename} already exists. Overwrite? (Y/N): ").strip().lower()
            if overwrite != 'y':
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                docx_filename = "{}_folder_structure_{}.docx".format(os.path.basename(folder_path), timestamp)

        doc.save(docx_filename)

    print(folder_structure)

if __name__ == "__main__":
    main()
